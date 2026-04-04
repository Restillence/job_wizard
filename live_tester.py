"""
MVP Live Evaluation — Arbeitsagentur + Arbeitnow Pipeline
Tests the actual MVP endpoints with real API data:
  1. POST /jobs/search-boards  (Arbeitsagentur + Arbeitnow)
  2. POST /resumes/upload
  3. POST /jobs/match
  4. POST /pipeline/search-and-match
Writes JSON output files for each endpoint.
"""
import io
import time
import json
import httpx
from docx import Document

from src.database import engine, SessionLocal
from src.models import Base, User

# ── Fresh DB ──────────────────────────────────────────────────────
Base.metadata.drop_all(bind=engine)
Base.metadata.create_all(bind=engine)

db = SessionLocal()
try:
    db.add(User(
        id="test_user_id",
        email="live@test.com",
        hashed_password="pwd",
        is_superuser=True,
    ))
    db.commit()
    print("[SEED] Created test user")
finally:
    db.close()

BASE_URL = "http://127.0.0.1:8000"


def write_out(name, data):
    with open(f"live_output_{name}.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"  [OK] Saved live_output_{name}.json")


def run():
    with httpx.Client(timeout=120) as client:
        # ── 1. Search Job Boards (Arbeitsagentur + Arbeitnow) ────
        print("\n═══ 1. POST /api/v1/jobs/search-boards ═══")
        search_payload = {
            "query": "Python Developer",
            "city": "Berlin",
            "country": "DE",
            "keywords": ["Python", "FastAPI"],
            "page": 1,
            "per_page": 10,
        }
        r = client.post(f"{BASE_URL}/api/v1/jobs/search-boards",
                        json=search_payload)
        print(f"  Status: {r.status_code}")
        try:
            out_boards = r.json()
            write_out("search_boards", out_boards)
            total = out_boards.get("total_found", 0)
            new = out_boards.get("newly_added", 0)
            print(f"  Found {total} jobs, {new} newly added")

            # Show first 3 jobs
            for j in out_boards.get("jobs", [])[:3]:
                print(f"    → {j['title']} @ {j['company_name']} ({j['source']})")
        except Exception as e:
            print(f"  [ERROR] {e} / body: {r.text[:300]}")
            return

        time.sleep(2)

        # ── 2. Resume Upload ──────────────────────────────────────
        print("\n═══ 2. POST /api/v1/resumes/upload ═══")
        doc = Document()
        doc.add_paragraph("Max Mustermann")
        doc.add_paragraph(
            "Senior Software Engineer aus München. "
            "Experte in Python, FastAPI, und Cloud-Infrastruktur. "
            "6 Jahre Erfahrung in agilen Teams. "
            "Schwerpunkte: Microservices, CI/CD, Kubernetes, PostgreSQL. "
            "Erfahrung mit Machine Learning und Datenanalyse."
        )
        buf = io.BytesIO()
        doc.save(buf)

        files = {
            "file": ("lebenslauf.docx", buf.getvalue(),
                     "application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document")
        }
        r = client.post(f"{BASE_URL}/api/v1/resumes/upload", files=files)
        print(f"  Status: {r.status_code}")
        write_out("resume_upload", r.json())

        time.sleep(2)

        # ── 3. Job Matching ───────────────────────────────────────
        print("\n═══ 3. POST /api/v1/jobs/match ═══")
        r = client.post(f"{BASE_URL}/api/v1/jobs/match",
                        json={"user_id": "test_user_id", "top_k": 10})
        print(f"  Status: {r.status_code}")
        out_match = r.json()
        write_out("jobs_match", out_match)
        matches = out_match.get("matched_jobs", [])
        print(f"  {len(matches)} matched jobs")
        for m in matches[:3]:
            print(f"    → {m['title']} @ {m['company_name']} "
                  f"(score: {m['similarity_score']})")

        time.sleep(2)

        # ── 4. Full Pipeline (board search + match, NO deep_search) ──
        print("\n═══ 4. POST /api/v1/pipeline/search-and-match ═══")
        pipeline_payload = {
            "cities": ["München"],
            "keywords": ["Backend", "Python"],
            "user_id": "test_user_id",
            "top_k": 10,
            "deep_search": False,
        }
        r = client.post(f"{BASE_URL}/api/v1/pipeline/search-and-match",
                        json=pipeline_payload)
        print(f"  Status: {r.status_code}")
        out_pipe = r.json()
        write_out("pipeline", out_pipe)
        print(f"  Board jobs found: {out_pipe.get('board_jobs_found', 0)}")
        print(f"  Board jobs new: {out_pipe.get('board_jobs_new', 0)}")
        for m in out_pipe.get("matched_jobs", [])[:3]:
            print(f"    → {m['title']} @ {m['company_name']} "
                  f"(score: {m['similarity_score']})")

        print("\n════════════════════════════════════════")
        print("[DONE] MVP Live Evaluation Complete!")
        print("════════════════════════════════════════")


if __name__ == "__main__":
    import threading
    import uvicorn
    from src.main import app

    def start_server():
        uvicorn.run(app, host="127.0.0.1", port=8000, log_level="warning")

    t = threading.Thread(target=start_server, daemon=True)
    t.start()
    time.sleep(4)
    run()
