import os
import io
from unittest.mock import patch, MagicMock
from docx import Document as DocxDocument


def test_upload_txt_resume(client):
    resume_content = "John Smith\nData Scientist\nEmail: john@email.com"

    with (
        patch("src.api.routers.resumes.pii_service.strip_pii") as mock_strip,
        patch("src.api.routers.resumes.generate_embedding") as mock_embedding,
    ):
        mock_strip.return_value = "[REDACTED]\nData Scientist\nEmail: [REDACTED]"
        mock_embedding.return_value = [0.1] * 3072

        response = client.post(
            "/api/v1/resumes/upload",
            files={"file": ("resume.txt", resume_content.encode(), "text/plain")},
        )

    assert response.status_code == 200
    data = response.json()
    assert "file_path" in data
    assert "resume_id" in data
    assert data["file_path"].endswith(".txt")

    file_path = data["file_path"]
    assert os.path.exists(file_path)

    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
        assert "[REDACTED]" in content

    if os.path.exists(file_path):
        os.remove(file_path)
    original = file_path.replace(".txt", ".txt")
    if os.path.exists(original):
        os.remove(original)


def test_upload_pdf_resume(client):
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("Jane Doe - Software Engineer")
    doc.add_paragraph("Python, Docker, Kubernetes")

    import fitz

    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text(
        (72, 72), "Jane Doe - Software Engineer\nPython, Docker, Kubernetes"
    )
    pdf_bytes = pdf_doc.tobytes()
    pdf_doc.close()

    with (
        patch("src.api.routers.resumes.pii_service.strip_pii") as mock_strip,
        patch("src.api.routers.resumes.generate_embedding") as mock_embedding,
    ):
        mock_strip.return_value = (
            "[REDACTED] - Software Engineer\nPython, Docker, Kubernetes"
        )
        mock_embedding.return_value = [0.1] * 3072

        response = client.post(
            "/api/v1/resumes/upload",
            files={"file": ("resume.pdf", pdf_bytes, "application/pdf")},
        )

    assert response.status_code == 200
    data = response.json()
    assert "resume_id" in data
    file_path = data["file_path"]
    assert file_path.endswith(".txt")

    for ext in [".txt", ".pdf"]:
        p = file_path.replace(".txt", ext)
        if os.path.exists(p):
            os.remove(p)


def test_upload_docx_resume(client):
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("Bob Builder")
    doc.add_paragraph("Project Manager")
    doc.save(buf)
    buf.seek(0)

    with (
        patch("src.api.routers.resumes.pii_service.strip_pii") as mock_strip,
        patch("src.api.routers.resumes.generate_embedding") as mock_embedding,
    ):
        mock_strip.return_value = "[REDACTED]\nProject Manager"
        mock_embedding.return_value = [0.1] * 3072

        response = client.post(
            "/api/v1/resumes/upload",
            files={
                "file": (
                    "resume.docx",
                    buf.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    data = response.json()
    file_path = data["file_path"]

    for ext in [".txt", ".docx"]:
        p = file_path.replace(".txt", ext)
        if os.path.exists(p):
            os.remove(p)


def test_upload_unsupported_format(client):
    response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.jpg", b"fake image data", "image/jpeg")},
    )

    assert response.status_code == 400
    assert "Unsupported" in response.json()["detail"]


def test_upload_empty_resume(client):
    with (
        patch("src.api.routers.resumes.pii_service.strip_pii") as mock_strip,
        patch("src.api.routers.resumes.generate_embedding") as mock_embedding,
    ):
        mock_strip.return_value = "   "
        mock_embedding.return_value = [0.1] * 3072

        response = client.post(
            "/api/v1/resumes/upload",
            files={"file": ("resume.txt", b"", "text/plain")},
        )

    assert response.status_code == 400
    assert "Could not extract" in response.json()["detail"]
