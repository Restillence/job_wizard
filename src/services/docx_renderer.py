import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.services.cv_parser import ParsedCV
from src.services.cv_generator import TailoredCV, CoverLetterResult


FONT_NAME = "Calibri"
FONT_COLOR = RGBColor(0x2D, 0x2D, 0x2D)
ACCENT_COLOR = RGBColor(0x1A, 0x56, 0xDB)
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
LINE_COLOR = RGBColor(0x1A, 0x56, 0xDB)


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10.5)
    font.color.rgb = FONT_COLOR

    for level in range(1, 4):
        heading_style_name = f"Heading {level}"
        if heading_style_name in doc.styles:
            heading = doc.styles[heading_style_name]
            heading.font.name = FONT_NAME
            heading.font.color.rgb = HEADING_COLOR


def _add_horizontal_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 80)
    run.font.color.rgb = LINE_COLOR
    run.font.size = Pt(6)


def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = FONT_NAME


def render_cv(tailored_cv: TailoredCV, parsed_cv: ParsedCV, output_path: str) -> str:
    doc = Document()

    sections = doc.sections
    if sections:
        section = sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _setup_styles(doc)

    name = parsed_cv.full_name or "[REDACTED]"
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = HEADING_COLOR
    name_run.font.name = FONT_NAME

    contact_parts = []
    if parsed_cv.email:
        contact_parts.append(parsed_cv.email)
    if parsed_cv.phone:
        contact_parts.append(parsed_cv.phone)
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(4)
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = FONT_COLOR

    _add_horizontal_line(doc)

    if tailored_cv.summary:
        _add_section_heading(doc, "Professional Summary")
        summary_p = doc.add_paragraph()
        summary_p.paragraph_format.space_after = Pt(8)
        summary_run = summary_p.add_run(tailored_cv.summary)
        summary_run.font.size = Pt(10.5)

    if tailored_cv.experience:
        _add_section_heading(doc, "Work Experience")
        for exp in tailored_cv.experience:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(2)
            title_run = title_p.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.name = FONT_NAME

            if exp.content:
                for line in exp.content.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    bullet_p = doc.add_paragraph(style="List Bullet")
                    bullet_p.paragraph_format.space_before = Pt(1)
                    bullet_p.paragraph_format.space_after = Pt(1)
                    clean = line.lstrip("•-* ").strip()
                    bullet_run = bullet_p.add_run(clean)
                    bullet_run.font.size = Pt(10)
                    bullet_run.font.name = FONT_NAME

    if tailored_cv.education:
        _add_section_heading(doc, "Education")
        for edu in tailored_cv.education:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(4)
            title_p.paragraph_format.space_after = Pt(2)
            title_run = title_p.add_run(edu.title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.name = FONT_NAME

            if edu.content:
                content_p = doc.add_paragraph()
                content_p.paragraph_format.space_after = Pt(4)
                content_run = content_p.add_run(edu.content)
                content_run.font.size = Pt(10)

    if tailored_cv.skills:
        _add_section_heading(doc, "Skills")
        skills_p = doc.add_paragraph()
        skills_p.paragraph_format.space_after = Pt(8)
        skills_run = skills_p.add_run(" | ".join(tailored_cv.skills))
        skills_run.font.size = Pt(10)

    if tailored_cv.languages:
        _add_section_heading(doc, "Languages")
        lang_p = doc.add_paragraph()
        lang_p.paragraph_format.space_after = Pt(8)
        lang_run = lang_p.add_run(", ".join(tailored_cv.languages))
        lang_run.font.size = Pt(10)

    if tailored_cv.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in tailored_cv.certifications:
            cert_p = doc.add_paragraph(style="List Bullet")
            cert_run = cert_p.add_run(cert)
            cert_run.font.size = Pt(10)

    if tailored_cv.additional_sections:
        for section in tailored_cv.additional_sections:
            _add_section_heading(doc, section.title)
            content_p = doc.add_paragraph()
            content_p.paragraph_format.space_after = Pt(8)
            content_run = content_p.add_run(section.content)
            content_run.font.size = Pt(10)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def render_cover_letter(
    cover_letter: CoverLetterResult,
    parsed_cv: ParsedCV,
    company_name: str,
    output_path: str,
) -> str:
    doc = Document()

    sections = doc.sections
    if sections:
        section = sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _setup_styles(doc)

    name = parsed_cv.full_name or "[REDACTED]"
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(2)
    header_run = header_p.add_run(name)
    header_run.bold = True
    header_run.font.size = Pt(14)
    header_run.font.name = FONT_NAME

    contact_parts = []
    if parsed_cv.email:
        contact_parts.append(parsed_cv.email)
    if parsed_cv.phone:
        contact_parts.append(parsed_cv.phone)
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(12)
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)

    _add_horizontal_line(doc)

    if company_name:
        company_p = doc.add_paragraph()
        company_p.paragraph_format.space_before = Pt(12)
        company_p.paragraph_format.space_after = Pt(4)
        company_run = company_p.add_run(company_name)
        company_run.font.size = Pt(10.5)

    salutation_p = doc.add_paragraph()
    salutation_p.paragraph_format.space_before = Pt(12)
    salutation_run = salutation_p.add_run("Dear Hiring Manager,")
    salutation_run.font.size = Pt(10.5)

    for paragraph_text in cover_letter.cover_letter.split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
        if paragraph_text.lower().startswith("dear"):
            continue

        body_p = doc.add_paragraph()
        body_p.paragraph_format.space_before = Pt(6)
        body_p.paragraph_format.space_after = Pt(6)
        body_run = body_p.add_run(paragraph_text)
        body_run.font.size = Pt(10.5)
        body_run.font.name = FONT_NAME

    closing_p = doc.add_paragraph()
    closing_p.paragraph_format.space_before = Pt(12)
    closing_lines = ["Sincerely,", "", name]
    closing_run = closing_p.add_run("\n".join(closing_lines))
    closing_run.font.size = Pt(10.5)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path
